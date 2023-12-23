import os
from pprint import pprint

import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE, WD_BUILTIN_STYLE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path
from pymorphy3 import MorphAnalyzer
from re import fullmatch


class WordFileEditor:

    def __init__(self, name_folder, file_name, text):
        """
        :param name_folder:
        :param file_name:
        :param text:
        """
        self.name_folder = name_folder
        self.file_name = file_name
        self.text = text
        self.relative_path = 'media/samples/your_doc_name.docx'
        self.current_dir_path = Path(__file__).resolve().parent.parent.parent
        self.file_path = os.path.join(self.current_dir_path, self.relative_path)

    # def main(self):
    #     document = docx.Document(self.file_path)
    #     text = "Н555555555555555555555ия."
    #     document.add_paragraph(text)
    #     document.save(self.file_path)

    def main(self):
        new_document = self.read_word(self.text)
        path_to_folder = self.create_folder(self.name_folder)
        self.write_word(new_document, path_to_folder, self.file_name)

    def read_word(self, text):
        document = docx.Document(self.file_path)
        style = document.styles.add_style('UserStyle', WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.font.underline = False
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        name_equipment = self.declension(self.name_folder)
        document.add_paragraph(
            f"\tДля предотвращения выхода из строя {name_equipment} прошу Вас приобрести:\n", style='UserStyle'
        )

        for style in document.styles:
            print("style.name == %s" % style.name)

        # styles = []
        # for paragraph in document.paragraphs:
        #     styles.append(paragraph.style)
        # pprint(styles)

        # times_new_roman_style = document.styles.add_style('Times New Roman', WD_STYLE_TYPE.CHARACTER)
        # times_new_roman_style.font.name = 'Times New Roman'

        for el in text:
            p = document.add_paragraph(
                el, style='List Number 3'
            )
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(14)


        return document

    def create_folder(self, name_folder):
        path = os.path.join(self.current_dir_path, f'media/orders/{name_folder}')
        if not os.path.isdir(path):
            os.mkdir(path)
            return path
        else:
            return path

    def declension(self, text):
        new_text = text.split()
        morph = MorphAnalyzer()
        gender = 'masc'
        for word in reversed(new_text):
            if self.check_word_is_rus(word) and len(word) > 3:
                gender = morph.parse(word)[0].tag.gender
                break

        for index, word in enumerate(new_text):
            if len(word) > 3:
                try:
                    parsed_word = morph.parse(word)[0]
                    if 'masc' in gender:
                        new_text[index] = parsed_word.inflect({'gent', 'masc'}).word
                    elif 'femn' in gender:
                        new_text[index] = parsed_word.inflect({'gent', 'femn'}).word
                except Exception as e:
                    if self.check_word_is_rus(word):
                        new_text[index] = word.lower()
        new_text = ' '.join(new_text)
        return new_text


    def check_word_is_rus(self, word):
        return bool(fullmatch(r'\b[а-яА-ЯЁё]+\b', word))


    def write_word(self, document, path_to_folder, file_name):
        full_path = rf'{path_to_folder}/{file_name}.docx'
        document.save(full_path)

